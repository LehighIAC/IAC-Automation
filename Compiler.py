"""
Fully-automated IAC report compiler
Usage: Copy all ARs into the ARs folder, Update info in Compiler.json5 and Utility.json5,
then run this script.
"""


import json5, os, locale, datetime
import pandas as pd
from easydict import EasyDict
from docx import Document, shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docxcompose.composer import Composer
from python_docx_replace import docx_replace, docx_blocks
from Shared.IAC import *

# If ARs/Sorted/ folder doesn't exist, create one
os.makedirs(os.path.join('ARs', 'Sorted'), exist_ok=True)

# If on macOS
if os.path.exists(os.path.join('Energy Charts', 'Energy Charts.fld')):
    chartPath = os.path.join('Energy Charts', 'Energy Charts.fld')
# If on Windows
elif os.path.exists(os.path.join('Energy Charts', 'Energy Charts_files')):
    chartPath = os.path.join('Energy Charts', 'Energy Charts_files')
else:
    # If chart html folder doesn't exist, exit
    raise Exception("Chart images not found. Please save the energy chart as web page (.htm).")

# Load config file and convert everything to local variables
print("Reading json5 database...", end ="")
jsonDict = json5.load(open('Compiler.json5'))
jsonDict.update(json5.load(open('Utility.json5')))
iac = EasyDict(jsonDict)
print("done")

# Initialize dataframe
columns = ["isAAR", "File Name", "ARC No.", "Description", "Electricity (kWh)", "Electricity (MMBtu)", "Demand (kW)"
           , "Natural Gas (MMBtu)", "Other Energy Type", "Other Energy Amount", "Other Resource Type", "Other Resource Amount"
           , "Savings Type", "Savings Value", "Annual Cost Savings", "Implementation Cost", "Payback Period"]
df = pd.DataFrame(columns=columns)

# Set locale to en_US
locale.setlocale(locale.LC_ALL, 'en_US')

print("Reading ARs...")
# Get all .docx files in ARs directory and extract information
ARList = [f for f in os.listdir('ARs') if f.endswith('.docx')]
AR_id = 0
for ARdoc in ARList:
    print(ARdoc)
    doc = Document(os.path.join('ARs', ARdoc))
    ARinfo = {}
    # check if the document is an AAR by title
    ARinfo['isAAR'] = ("AAR" in doc.paragraphs[0].text.split(':')[0])
    # Record file name
    ARinfo['File Name'] = ARdoc
    # Parse the title of the .docx file
    ARinfo['Description'] = doc.paragraphs[0].text.split(':')[1].strip()
    # Read the 1st table in .docx files
    try:
        table = doc.tables[0]
    except:
        print("Error: " + ARdoc + " is not a valid AR. Please check if the summary table is present.")
        exit()
    for row in table.rows:
        key = row.cells[0].text
        value = row.cells[1].text
        # Parse ARC Number
        if "arc" in key.lower() and "number" in key.lower():
            validate_arc(value)
            ARinfo['ARC No.'] = value
        # Parse Annual Cost Savings
        elif "annual" in key.lower() and "cost" in key.lower():
            # convert currency to interger
            ARinfo['Annual Cost Savings'] = locale.atoi(value.strip("$"))
        # Parse Implementation Cost
        elif "implementation" in key.lower():
            # convert currency to interger
            ARinfo['Implementation Cost'] = locale.atoi(value.strip("$"))
        # If Payback Period skip (Doesn't matter, will calculate later)
        elif "payback" in key.lower():
            continue
        # Parse Electricity
        elif "electricity" in key.lower():
            ARinfo['Electricity (kWh)'] = locale.atoi(value.split(' ')[0])
        # Parse Demand
        elif "demand" in key.lower():
            ARinfo['Demand (kW)'] = locale.atoi(value.split(' ')[0])
        # Parse Natural Gas
        elif "natural" in key.lower():
            ARinfo['Natural Gas (MMBtu)'] = locale.atoi(value.split(' ')[0])
        # Parse undefined type
        else:
            # If the value contains mmbtu, parse it as other energy
            if "mmbtu" in value.lower():
                # Remove the last word (usually "savings")
                ARinfo['Other Energy Type'] = key.rsplit(' ', 1)[0]
                # Parse number
                ARinfo['Other Energy Amount'] = locale.atoi(value.split(' ')[0])
            # If not, parse it as other resource
            else:
                # Remove the last word (usually "savings")
                ARinfo['Other Resource Type'] = key.rsplit(' ', 1)[0]
                # Keep the whole string
                ARinfo['Other Resource Amount'] = value   
    # Add dictionary to dataframe
    for key in ARinfo:
        df.loc[AR_id, key] = ARinfo[key]
    AR_id += 1
print("done")

print("Analyzing ARs...", end ="")
## Calculate on columns
# Calculate payback period
df['Payback Period'] = df['Implementation Cost'] / df['Annual Cost Savings']
# Convert electricity to MMBtu
df['Electricity (MMBtu)'] = df['Electricity (kWh)']* 0.003413/0.33
# Sort df by payback period
df = df.sort_values(by=['Payback Period'])

## Format energy savings strings
for index, row in df.iterrows():
    ST = ""
    SV = ""
    if pd.notna(row['Electricity (kWh)']):
        ST = ST + "Electricity" + '\n\n'
        SV = SV + locale.format_string('%d',row['Electricity (kWh)'], grouping=True) + ' kWh' + '\n'
        SV = SV + '(' + locale.format_string('%d',row['Electricity (MMBtu)'], grouping=True) + ' MMBtu)' + '\n'
    if pd.notna(row['Demand (kW)']):
        ST = ST + "Demand" + '\n'
        SV = SV + locale.format_string('%d',row['Demand (kW)'], grouping=True) + ' kW' + '\n'
    if pd.notna(row['Natural Gas (MMBtu)']):
        ST = ST + "Natural Gas" + '\n'
        SV = SV + locale.format_string('%d',row['Natural Gas (MMBtu)'], grouping=True)  + ' MMBtu' + '\n'
    if pd.notna(row['Other Energy Type']):
        ST = ST + row['Other Energy Type'] + '\n'
        SV = SV + locale.format_string('%d',row['Other Energy Amount'], grouping=True)  + ' MMBtu' '\n'
    if pd.notna(row['Other Resource Type']):
        ST = ST + row['Other Resource Type'] + '\n'
        SV = SV + row['Other Resource Amount'] + '\n'
    ST = ST.rstrip('\n')
    SV = SV.rstrip('\n')
    df.at[index, 'Savings Type'] = ST
    df.at[index, 'Savings Value'] = SV

## Summation statistics
# Filter ARs
AR_df = df[df['isAAR'] == False]
# Reorder index
AR_df = AR_df.reset_index(drop=True)
# AR statistics
EkWh = AR_df['Electricity (kWh)'].sum(axis=0, skipna=True)
EMMBtu = AR_df['Electricity (MMBtu)'].sum(axis=0, skipna=True)
NMMBtu = AR_df['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
OMMBtu = AR_df['Other Energy Amount'].sum(axis=0, skipna=True)
# Add up all energy in MMBtu
iac.ARMMBtu = round(EMMBtu + NMMBtu + OMMBtu)
# Calculate CO2 (Currently other type of energy ignored)
iac.CO2 = round((53 * NMMBtu + 0.22 * EkWh)/1000)
# Add up all cost
iac.ARACS = AR_df['Annual Cost Savings'].sum(axis=0, skipna=True)
iac.ARIC = AR_df['Implementation Cost'].sum(axis=0, skipna=True)
# Payback period in number
iac.ARPB = round(iac.ARIC / iac.ARACS, 1)
# Payback period in formatted string
iac.PB = payback(iac.ARACS, iac.ARIC)
print("done")

print("Reformatting ARs...", end ="")
## Reformatting ARs
for index, row in AR_df.iterrows():
    doc = Document(os.path.join('ARs', row['File Name']))
    # Change title and make it upper case
    doc.paragraphs[0].text = "AR "+ str(index+1) + ': ' + row['Description'].upper()
    # Set font to bold
    doc.paragraphs[0].runs[0].bold = True
    # Set font to upright
    doc.paragraphs[0].runs[0].italic = False
    # Set font to no underline
    doc.paragraphs[0].runs[0].underline = False
    # Set font size to 12
    doc.paragraphs[0].runs[0].font.size = shared.Pt(12)
    # Add pagebreak to the end of the document
    doc.add_page_break()
    doc.save(os.path.join('ARs', 'Sorted', 'AR'+ str(index+1) + '.docx'))
print("done")

# Check if there's at least 1 AAR
AAR = df['isAAR'].any()
if AAR:
    print("Analyzing AARs...", end ="")
    # Filter AAR
    AAR_df = df[df['isAAR'] == True]
    # reorder index
    AAR_df = AAR_df.reset_index(drop=True)
    # AAR statistics
    EMMBtu = AAR_df['Electricity (MMBtu)'].sum(axis=0, skipna=True)
    NMMBtu = AAR_df['Natural Gas (MMBtu)'].sum(axis=0, skipna=True)
    OMMBtu = AAR_df['Other Energy Amount'].sum(axis=0, skipna=True)
    # Add up all energy
    iac.AARMMBtu = round(EMMBtu + NMMBtu + OMMBtu)
    # Add up all cost
    iac.AARACS = AAR_df['Annual Cost Savings'].sum(axis=0, skipna=True)
    iac.AARIC = AAR_df['Implementation Cost'].sum(axis=0, skipna=True)
    # Payback period in number
    iac.AARPB = round(iac.AARIC / iac.AARACS, 1)
    print("done")

    print("Reformatting AARs...", end ="")
    # Modify the title of the AAR docx
    for index, row in AAR_df.iterrows():
        doc = Document(os.path.join('ARs', row['File Name']))
        # Change title and make it upper case
        doc.paragraphs[0].text = "AAR "+ str(index+1) + ': ' + row['Description'].upper()
        # Set font to bold
        doc.paragraphs[0].runs[0].bold = True
        # Set font to upright
        doc.paragraphs[0].runs[0].italic = False
        # Set font to no underline
        doc.paragraphs[0].runs[0].underline = False
        # Set font size to 12
        doc.paragraphs[0].runs[0].font.size = shared.Pt(12)
        # Add pagebreak to the end of the document
        doc.add_page_break()
        doc.save(os.path.join('ARs', 'Sorted', 'AAR'+ str(index+1) + '.docx'))
    print("done")

print("Parsing plant information...", end ="")
## Compiler.json5 Calculations
# Report date = today or 60 days after assessment, which ever is earlier
VD = datetime.datetime.strptime(iac.VDATE, '%B %d, %Y')
RDATE = min(datetime.datetime.today(), VD + datetime.timedelta(days=60))
iac.RDATE = datetime.datetime.strftime(RDATE, '%B %d, %Y')
# Sort participant and contributor name list
PART=""
for name in iac.PARTlist:
    # Sort by last name
    iac.PARTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
    PART  = PART + name + '\n'
iac.PART = PART.rstrip('\n')
iac.pop('PARTlist')
CONT=""
for name in iac.CONTlist:
    # Sort by last name
    iac.CONTlist.sort(key=lambda x: x.rsplit(' ', 1)[1])
    CONT  = CONT + name + '\n'
iac.CONT = CONT.rstrip('\n')
iac.pop('CONTlist')
print("done")

## Format strings
# set electricity cost to 3 digits accuracy
iac = dollar(['EC'],iac,3)
# set the natural gas and demand to 2 digits accuracy
iac = dollar(['DC', 'FC'],iac,2)
# set the rest to integer
varList = ['ARACS', 'ARIC', 'AARACS', 'AARIC', 'TotalECost', 'TotalFCost', 'TotalCost']
iac = dollar(varList,iac,0)
# Format all numbers to string with thousand separator
iac = grouping_num(iac)

## Load introduction template
doc_intro = Document(os.path.join('Report', 'Introduction.docx'))

# Add rows to AR table (Should be the 3rd table)
print("Writing AR table...", end ="")
locale._override_localeconv={'frac_digits':0}
ARTable = doc_intro.tables[2]
for index, row in AR_df.iterrows():
    ARrow = ARTable.rows[index+1].cells
    # Add ARC No.
    ARrow[0].text = 'AR ' + str(index+1) + '\n' + row['ARC No.']
    ARrow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add description
    ARrow[1].text = row['Description']
    ARrow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add savings type
    ARrow[2].text = row['Savings Type']
    ARrow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add savings value
    ARrow[3].text = row['Savings Value']
    ARrow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add annual cost savings
    ARrow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
    ARrow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add implementation cost
    ARrow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
    ARrow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add payback period
    ARrow[6].text = str(round(row['Payback Period'],1))
    ARrow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set 3pt before and after paragraph
    for col in range(0,7):
        ARrow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
        ARrow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
# Delete unused rows (Currectly row 1-15 are empty)
for index in reversed(range(len(AR_df), 15)):
    ARTable._tbl.remove(ARTable.rows[index+1]._tr)
print("done")

if AAR:
    # Add rows to AAR table (Should be the 4th table)
    print("Writing AAR table...", end ="")
    locale._override_localeconv={'frac_digits':0}
    AARTable = doc_intro.tables[3]
    for index, row in AAR_df.iterrows():
        AARrow = AARTable.rows[index+1].cells
        # Add ARC No.
        AARrow[0].text = 'AAR ' + str(index+1) + '\n' + row['ARC No.']
        AARrow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add description
        AARrow[1].text = row['Description']
        AARrow[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add savings type
        AARrow[2].text = row['Savings Type']
        AARrow[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add savings value
        AARrow[3].text = row['Savings Value']
        AARrow[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add annual cost savings
        AARrow[4].text = locale.currency(row['Annual Cost Savings'], grouping=True)
        AARrow[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add implementation cost
        AARrow[5].text = locale.currency(row['Implementation Cost'], grouping=True)
        AARrow[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Add payback period
        AARrow[6].text = str(round(row['Payback Period'],1))
        AARrow[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Set 3pt before and after paragraph
    for col in range(0,7):
        AARrow[col].paragraphs[0].paragraph_format.space_before = shared.Pt(3)
        AARrow[col].paragraphs[0].paragraph_format.space_after = shared.Pt(3)
    # Delete unused rows (Currectly row 1-5 are empty)
    for index in reversed(range(len(AAR_df), 5)):
        AARTable._tbl.remove(AARTable.rows[index+1]._tr)
    print("done")
else:
    # delete this table
    doc_intro._body._body.remove(doc_intro.tables[3]._tbl)

# Remove AAR blocks if no AAR
docx_blocks(doc_intro, AAR = AAR)

# Replacing keys
print("Replacing keys in introduction...", end ="")
docx_replace(doc_intro, **iac)
print("done")

# Save introduction
filename_intro = iac.LE + '-intro.docx'
doc_intro.save(filename_intro)

## Load backgroud template
doc_back = Document(os.path.join('Report', 'Background.docx'))

# Replacing keys
print("Replacing keys in background...", end ="")
docx_replace(doc_back, **iac)
print("done")

# Save background
filename_back = iac.LE + '-back.docx'
doc_back.save(filename_back)

## Load energy bill analysis template
doc_energy = Document(os.path.join('Report', 'Energy.docx'))

# Add energy chart images
print("Adding energy chart images...", end ="")
# If on macOS
if chartPath == os.path.join('Energy Charts', 'Energy Charts.fld'):
    add_image(doc_energy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
    add_image(doc_energy, '#ECChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
    add_image(doc_energy, '#DUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
    add_image(doc_energy, '#DCChart', os.path.join(chartPath, "image004.png"), shared.Inches(6))
    add_image(doc_energy, '#FUChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
    add_image(doc_energy, '#FCChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
    add_image(doc_energy, '#PieUChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
    add_image(doc_energy, '#PieCChart', os.path.join(chartPath, "image008.png"), shared.Inches(6))
    add_image(doc_energy, '#TotalChart', os.path.join(chartPath, "image009.png"), shared.Inches(9))
# If on Windows
elif chartPath == os.path.join('Energy Charts', 'Energy Charts_files'):
    add_image(doc_energy, '#EUChart', os.path.join(chartPath, "image001.png"), shared.Inches(6))
    add_image(doc_energy, '#ECChart', os.path.join(chartPath, "image002.png"), shared.Inches(6))
    add_image(doc_energy, '#DUChart', os.path.join(chartPath, "image003.png"), shared.Inches(6))
    add_image(doc_energy, '#DCChart', os.path.join(chartPath, "image005.png"), shared.Inches(6))
    add_image(doc_energy, '#FUChart', os.path.join(chartPath, "image006.png"), shared.Inches(6))
    add_image(doc_energy, '#FCChart', os.path.join(chartPath, "image007.png"), shared.Inches(6))
    add_image(doc_energy, '#PieUChart', os.path.join(chartPath, "image009.png"), shared.Inches(6))
    add_image(doc_energy, '#PieCChart', os.path.join(chartPath, "image011.png"), shared.Inches(6))
    add_image(doc_energy, '#TotalChart', os.path.join(chartPath, "image013.png"), shared.Inches(9))
print("done")

# Fill in energy chart tables from Energy Charts.xlsx
print("Adding energy chart tables...", end ="")
# Read electricity table from B6 to I19
edf = pd.read_excel(os.path.join('Energy Charts', 'Energy Charts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'B:I')
# Read fuel table from K6 to N19
fdf = pd.read_excel(os.path.join('Energy Charts', 'Energy Charts.xlsx'), sheet_name="Raw Data", skiprows = 5, nrows=13, usecols = 'K:N')

# Add rows to electricity table (Should be the 1st table)
etable = doc_energy.tables[0]
for index, row in edf.iterrows():
    erow = etable.rows[index+3].cells
    # Add Month
    erow[0].text = edf.iloc[(index, 0)]
    erow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in range(1,8):
        # Add interger with thousand separator
        erow[col].text = locale.format_string('%d',round(edf.iloc[(index, col)]), grouping=True)
        erow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Bold the last row
    if index == 12:
        for col in range(0,8):
            erow[col].paragraphs[0].runs[0].bold = True

# Add rows to fuel table (Should be the 2nd table)
ftable = doc_energy.tables[1]
for index, row in fdf.iterrows():
    frow = ftable.rows[index+3].cells
    # Add Month
    frow[0].text = fdf.iloc[(index, 0)]
    frow[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in range(1,4):
        # Add interger with thousand separator
        frow[col].text = locale.format_string('%d',round(fdf.iloc[(index, col)]), grouping=True)
        frow[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Bold the last row
    if index == 12:
        for col in range(0,4):
            frow[col].paragraphs[0].runs[0].bold = True
print("done")
# Replacing keys
print("Replacing keys in energy charts...", end ="")
docx_replace(doc_energy, **iac)
print("done")
# Save energy charts
filename_energy = iac.LE + '-energy.docx'
doc_energy.save(filename_energy)

print("Combining all docs...", end ="")
# A list of docs to combine
docList = [os.path.join('Report', 'ToC.docx')]
ARList = []
for ARlen in range(1, len(AR_df)+1):
    ARList.append(os.path.join('ARs', 'Sorted','AR' + str(ARlen) + '.docx'))
docList.extend(ARList)
if AAR:
    docList.append(os.path.join('Report', 'AAR.docx'))
    AARList = []
    for AARlen in range(1, len(AAR_df)+1):
        AARList.append(os.path.join('ARs', 'Sorted','AAR' + str(AARlen) + '.docx'))
    docList.extend(AARList)
else:
    pass
# docList.append(os.path.join('Report','Background.docx'))
docList.append(filename_back)
# docList.append(os.path.join('Report','Background.docx'))
docList.append(filename_energy)

# Combine all docx files
master = Document(filename_intro)
composer = Composer(master)
for i in range(0, len(docList)):
    doc_temp = Document(docList[i])
    composer.append(doc_temp)
filename = iac.LE +'.docx'
composer.save(filename)
# delete temp files
os.remove(filename_intro)
os.remove(filename_back)
os.remove(filename_energy)
print("done")

# Open the combined docx file
doc = Document(filename)
# Change the orientation of the last section back to landscape
# If anything goes wrong, check if there's a section break
# at the end of background.docx
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Save final report
doc.save(filename)
print(filename + " is finished.")

# Caveats
caveat("Please add Process Description, Major Equipment, Current Best Practices, and plant layout image.")
caveat("Please refresh ToC, tables and figures after running this script.")